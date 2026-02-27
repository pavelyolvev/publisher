import os
import sys
import re

import doc2docx
from docx import Document
from datetime import date

from html import table_gen

template_p1 = '<p style="text-align: center;"><strong>Постановление </strong></p>\n<p style="text-align: right;"><strong>{}</strong></p>\n<p style="text-align: center;"><strong>{}</strong></p>\n'
template_p2 = '<p style="text-align: justify;">{}</p>\n'
template_p3 = '<p style="text-align: center;"><strong>П О С Т А Н О В Л Я Е Т:</strong></p>\n'
template_p3_short_offline = '<p style="text-align: center;"><strong>П О С Т А Н О В Л Я Е Т: <a href="http://pohr.ru/wps/wp-content/uploads/{}/{}/{}">см. приложение</a></strong></p>\n'
template_p3_short = '<p style="text-align: center;"><strong>П О С Т А Н О В Л Я Е Т: <a href="{}">см. приложение</a></strong></p>\n'
template_p4 = '<p style="text-align: justify;">{}</p>\n'
template_p5_offline = '<p style="text-align: center;">{}   <strong>{}</strong></p>\n<p style="text-align: right;"><strong><a href="http://pohr.ru/wps/wp-content/uploads/{}/{}/{}">Приложение</a></strong></p>\n'
template_p5 = '<p style="text-align: center;">{}   <strong>{}</strong></p>\n<p style="text-align: right;"><strong><a href="{}">Приложение</a></strong></p>\n'
# для pohr ru должно быть http://pohr.ru/wps/wp-content/....
text_limiter_number = 2750  # количество разрешенных символов после заголовка "постановление:".


# если больше - генерируется по template_p3_short
def time_limiter(num):
    res_h = 0
    res_m = 0
    if num >= 1440:
        num = num % 1440
    res_h = num // 60
    res_m = num % 60
    print(res_h, res_m)
    return res_h, res_m


def parse_document_text_table(filepath, is_online):
    extension = os.path.splitext(filepath)[1].lstrip(".")

    if extension == "doc":
        # Создаем имя для конвертированного файла
        base_name = os.path.basename(filepath)
        docx_name = base_name + "x"  # добавляем 'x' к имени файла

        # Создаем директорию для конвертированных файлов, если её нет
        converted_dir = os.path.join(os.path.dirname(filepath), "convertedDoc")
        os.makedirs(converted_dir, exist_ok=True)

        # Полный путь для сохранения
        output_path = os.path.join(converted_dir, docx_name)

        # Конвертируем
        try:
            doc2docx.convert(filepath, output_path)
            path = output_path  # Используем конвертированный файл
        except Exception as e:
            print(f"Ошибка конвертации: {e}")
            return None, None, None
    else:
        path = filepath  # Используем исходный файл
    print(is_online)
    num_lvl_counter = [0] * 9
    global doc_date, header
    doc_num = 0
    html = []
    is_before = True
    is_empty = True
    text_before_postanovlenie = []
    text_after_postanovlenie = []
    name_html = ""
    tables = []
    table_html = ""
    author_html = ""
    if path != "":
        f = open(f'{path}', 'rb')
        document = Document(f)
        lowest_num_lvl = 0
        for i, block in enumerate(document.element.body):
            # Получаем локальное имя тега (без пространства имён)
            tag = block.tag.split('}')[-1]  # отрезаем {namespace}
            to_append = ""
            if tag == 'tbl':
                # print(f"\n[{i}] 🔲 НАЧАЛО ТАБЛИЦЫ")

                # Получаем объект таблицы
                table = document.tables[len([b for b in document.element.body[:i] if b.tag.endswith('tbl')])]
                # print(f"    Размер: {len(table.rows)}x{len(table.columns)}")
                docdate, header = is_table_name(table)
                if docdate is not None:
                    doc_date = docdate
                    doc_num = int(doc_date.split('№')[1])
                else:
                    #to_append = gen_html_table(table)
                     to_append = table_gen.gen_html_table_simple(table)
                # print(doc_date)

            elif tag == 'p':
                paragraph = document.paragraphs[len([b for b in document.element.body[:i] if b.tag.endswith('p')])]
                prev_paragraph = document.paragraphs[
                    len([b for b in document.element.body[:i] if b.tag.endswith('p')]) - 1]
                # if paragraph.text.strip():
                # print(''.join(paragraph.text.split()) != "")
                # print(is_empty)
                if ''.join(paragraph.text.split()) != "" and is_empty:
                    is_empty = False
                    text_before_postanovlenie.clear()

                to_append = ' '.join(paragraph.text.split())
                is_numbered = paragraph._element.pPr is not None and paragraph._element.pPr.numPr is not None
                num_pattern = r'\s*(\d+\.)+\s*'
                ilvl = -1
                match_in_ppr = False
                number_in_list = re.match(num_pattern, paragraph.text)
                if number_in_list:
                    # print(f"re.findall = {re.findall(r'\d+\.', number_in_list.group())}")
                    ilvl = len(re.findall(r'\d+\.', number_in_list.group())) - 1
                    match_in_ppr = True
                if (is_numbered):
                    ilvl = paragraph._element.pPr.numPr.ilvl.val if paragraph._element.pPr.numPr.ilvl is not None else 0
                if ilvl != -1:
                    num_lvl_counter[ilvl] += 1
                    lowest_num_lvl = ilvl if ilvl > lowest_num_lvl else lowest_num_lvl
                    if ilvl < lowest_num_lvl:
                        for i in range(ilvl + 1, len(num_lvl_counter)):
                            num_lvl_counter[i] = 0
                    number = ".".join([str(i) for i in num_lvl_counter if i != 0])
                    if not match_in_ppr:
                        to_append = number + ". " + to_append
                try:
                    author_par, author = is_author(prev_paragraph.text, to_append)
                    if author_par:
                        today = date.today()
                        # author_html = f'<p style="text-align: center;">{author_par.replace(author, "")}   <strong>{author}</strong></p>'
                        if is_online:
                            author_html = template_p5.format(author_par, author, '{}')
                        else: author_html = template_p5_offline.format(author_par, author, today.strftime('%Y'), today.strftime('%m'), os.path.basename(filepath))
                        break
                except TypeError:
                    print("no author here")

                # print(f"\n[{i}] 📝 Параграф: {paragraph.text}")
            elif tag == 'sectPr':
                print(f"\n[{i}] 📄 Секция (конец документа/раздела)")

            if is_post(to_append) and is_before:
                is_before = False
                continue
            if is_before:
                text_before_postanovlenie.append(to_append)
            else:
                text_after_postanovlenie.append(to_append)

        print(text_before_postanovlenie)
        print(text_after_postanovlenie)
        today = date.today()
        html.append(gen_name_before_html(doc_date, text_before_postanovlenie, header))
        if len(''.join(text_after_postanovlenie)) > text_limiter_number:
            if is_online:
                html.append(template_p3_short.format("{}"))
            else: html.append(template_p3_short_offline.format(today.strftime('%Y'), today.strftime('%m'), os.path.basename(filepath)))
            print(''.join(html))
        else:
            print(doc_date)
            html.append(template_p3)
            html.append(gen_text_after_html(text_after_postanovlenie))
            html.append(author_html)
        # print(''.join(html))
        return ''.join(html), doc_date, doc_num


def is_table_name(table):
    matched = None # номер
    non_default_matched = None # если не None - заголовок
    if table.rows:
        for i, t_row in enumerate(table.rows):
            for j, cell in enumerate(table.rows[i].cells):
                print(f"строка {i}, ячейка {j}: {cell.text}")
                match = re.match(r'\s*администрация.*постановление.*(\d{2}\.\d{2}\.\d{4}\s*№\s*\d+)',
                                 ' '.join(cell.text.lower().split()))
                if match:
                    matched = match.group(1)
                elif matched is not None:
                    non_default_matched = cell.text
    return matched, non_default_matched


def gen_html_table(table):
    cell_template = "<td style='border-width:1px; border-style:solid;'>{}</td>"
    row_template = "<tr>\n{}\n</tr>"

    cells = []
    rows = []

    for i, row in enumerate(table.rows):
        for cell in table.rows[i].cells:
            cells.append(cell_template.format(' '.join(cell.text.split())))
        rows.append(row_template.format('\n'.join(cells)))
    table_html = f"<table style='border-collapse: collapse'>\n{''.join(rows)}\n</table>"
    return table_html


def colspan_for_cell(c, cur_row, table): # c - индекс ячейки в строке
                                         # cur_row - индекс текущей строки
    colspan = 1
    print("    vert:")
    cells = table.rows[c].cells
    print(cells[cur_row])
    for j in range(cur_row, len(cells)):
        if j + 1 in range(0, len(cells)) and cells[j + 1] == cells[j]:
            for i in range(j, len(cells)):
                print(cells[i])
                print('    end vert')
                if i + 1 in range(0, len(cells)) and cells[i + 1] == cells[i]:
                    colspan += 1
    print(f"   colspan: {colspan}")
    return colspan


def gen_name_before_html(doc_date_, text_before, header_):
    html = []
    name = ""
    if header_ is None:
        name_arr = []
        for i, line in enumerate(text_before):
            if line == "":
                name_arr = text_before[0:i:1]
                print(text_before[0:i])
                del text_before[0:i]
                break
        name = ' '.join(name_arr)
    else: name = header_
    html.append(template_p1.format(doc_date_, name))

    for line in text_before:
        if line != "":
            html.append(template_p2.format(line))
    return ''.join(html)


def gen_text_after_html(text_after):
    html = []
    for i, line in enumerate(text_after):
        if line == "":
            continue
        else:
            match = re.search(r'(<table>.*?</table>)', '\n'.join(line))
            if match:
                html.append(line)
            else:
                html.append(template_p2.format(line))
    return ''.join(html)

    """parts = re.split(r'(<table>.*?</table>)', '\n'.join(text_after), flags=re.DOTALL)
    
    for i, part in enumerate(parts):
        if part.strip():  # пропускаем пустые части
            if part.startswith('<table>'):
                html.append(part + "\n")
            else:
                html.append(template_p2.format(part.strip()))
    return ''.join(html)"""


def is_post(paragraph) -> bool:
    if ''.join(paragraph.lower().split()) == "ПОСТАНОВЛЯЕТ:".lower():
        return True


def is_author(prev_paragraph, paragraph):
    name_pattern = r'([А-ЯЁ]\.\s*[А-ЯЁ]\.\s*[А-ЯЁ][а-яё]+)'
    match = re.search(name_pattern, paragraph)
    if prev_paragraph == "" and match and len(paragraph) < 150:
        author = re.split(name_pattern, paragraph)
        return ' '.join(author[0].split()), author[1]


if len(sys.argv) != 2:
    print("Использование: python script.py <input_file> <output_file>")
    # sys.exit(1)
else:
    input_file = sys.argv[1]
    html, datet, num = parse_document_text_table(input_file)
    print(html)

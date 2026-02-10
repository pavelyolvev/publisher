import os
from tkinter import *
from tkinter import ttk, filedialog, scrolledtext, messagebox
import re
from docx import Document
import pandas as pd
from datetime import date

counter = 0
num_lvl_counter = [0] * 9
filepath = ""
author = ""
tables = []

doc_date = ""

template_p1 = '<p style="text-align: center;"><strong>Постановление </strong></p>\n<p style="text-align: right;"><strong>{}</strong></p>\n<p style="text-align: center;"><strong>{}</strong></p>\n'
template_p2 = '<p style="text-align: justify;">{}</p>\n'
template_p3 = '<p style="text-align: center;"><strong>П О С Т А Н О В Л Я Е Т:</strong></p>\n'
template_p4 = '<p style="text-align: justify;">{}</p>\n'
template_p5 = '<p style="text-align: center;">{}   <strong>{}</strong></p>\n<p style="text-align: right;"><strong><a href="http://pohr.ru/wps/wp-content/uploads/{}/{}/{}.docx">Приложение</a></strong></p>\n'

full_text = []
name = []
text_before_postanovlenie = []
text_after_postanovlenie = []

def open_file():
    global filepath
    filepath = filedialog.askopenfilename()
    if filepath != "":
        if check_if_doc:
            lbl["text"] = f"File Path: {filepath}"
            #read_doc()
def check_if_doc(): # заглушка
    return True
        

def read_doc():
    global filepath
    global tables
    if filepath != "":
        f = open(f'{filepath}', 'rb')
        document = Document(f)
        all_tables = document.tables
        print('Всего таблиц в документе:', len(all_tables))
        data_tables = {i:None for i in range(len(all_tables))}
        # проходимся по таблицам
        for i, table in enumerate(all_tables):
            #print('\nДанные таблицы №', i)
            # создаем список строк для таблицы `i` (пока пустые)
            data_tables[i] = [[] for _ in range(len(table.rows))]
            # проходимся по строкам таблицы `i`
            for j, row in enumerate(table.rows):
                # проходимся по ячейкам таблицы `i` и строки `j`
                for cell in row.cells:
                    # добавляем значение ячейки в соответствующий
                    # список, созданного словаря под данные таблиц
                    data_tables[i][j].append(cell.text)
                    data_pattern = r'\d{2}\.\d{2}\.\d{4}\s*№\s*\d+'
                    match = re.search(data_pattern, cell.text)
                    if match:
                        print(match.group())
                        global doc_date
                        if doc_date == "":
                            doc_date = match.group()
                        break

            # смотрим извлеченные данные 
            # (по строкам) для таблицы `i`
            #print(data_tables[i])
            #print('\n')

        #print('Данные всех таблиц документа:')
        #print(data_tables)
        extracting = True
        lowest_num_lvl = 0
        for i, paragraph in enumerate(document.paragraphs):
            print(paragraph.text + "\n" + paragraph.style.name)
            style_name = paragraph.style.name if paragraph.style else "Normal"
            is_numbered = paragraph._element.pPr is not None and paragraph._element.pPr.numPr is not None
            num_pattern = r'\s*(\d+\.)+\s*'
            ilvl = -1
            match_in_ppr = False
            number_in_list = re.match(num_pattern, paragraph.text)
            if number_in_list:
                print(f"re.findall = {re.findall(r'\d+\.', number_in_list.group())}")
                ilvl = len(re.findall(r'\d+\.', number_in_list.group()))-1
                match_in_ppr = True
            if (is_numbered):
                ilvl = paragraph._element.pPr.numPr.ilvl.val if paragraph._element.pPr.numPr.ilvl is not None else 0
                #prefixes = ["1.", "a."]
                #prefix = prefixes[min(ilvl, len(prefixes)-1)]
            if ilvl != -1:                
                num_lvl_counter[ilvl] += 1
                lowest_num_lvl = ilvl if ilvl > lowest_num_lvl else lowest_num_lvl
                if ilvl < lowest_num_lvl:
                    for i in range(ilvl+1, len(num_lvl_counter)):
                        num_lvl_counter[i] = 0
                number = ".".join([str(i) for i in num_lvl_counter if i != 0])
                if match_in_ppr: full_text.append(paragraph.text.strip())
                else: full_text.append(number + ". " + paragraph.text.strip())
                print(f"Lvl: {ilvl}\nFinal number: {number}")
            else:    
                full_text.append(paragraph.text.strip())
            # Вырезка сведений о постановлении
            if extracting == True:
                if paragraph.text != "" :
                    text = paragraph.text
                    name.append(' '.join(text.split()))
                if name != [] and paragraph.text == "": extracting = False
                else: continue
                
        print(' '.join(name))
        print(name[len(name)-1])
       
        start_before_ind = full_text.index(name[len(name)-1]) if len(name) > 1 else full_text.index(name[len(name)])
        start_after_ind = 0
        name_pattern = r'([А-ЯЁ]\.\s*[А-ЯЁ]\.\s*[А-ЯЁ][а-яё]+)'
        for i in range(start_before_ind+1, len(full_text)):
            if paragraph.text != "" :
                continue
            else:
                if ''.join(full_text[i].lower().split()) != "ПОСТАНОВЛЯЕТ:".lower():
                    text_before_postanovlenie.append(' '.join(full_text[i].split()))
                else: 
                    start_after_ind = i
                    break
        for i in range(start_after_ind+1, len(full_text)):
            if paragraph.text != "" :
                continue
            else:
                if full_text[i-1] == "" and re.search(name_pattern, full_text[i]) and len(full_text[i]) < 150: #Если находит строку с фамилией автора в коротком абзаце
                    global author
                    author = re.split(name_pattern ,full_text[i])
                    print(author)
                    break
                else: text_after_postanovlenie.append(' '.join(full_text[i].split()))
        print("\n")
        #print(' '.join(text_before_postanovlenie))
        print(text_before_postanovlenie)
        print("\n")
        #print(' '.join(text_after_postanovlenie))
        
        f.close()
        print("\n")
        #print(full_text)
        save_to_file()
    return
    
def save_to_file():
    content = []
    content.append(template_p1.format(doc_date, ' '.join(name)))
    for paragraph in text_before_postanovlenie:
        if paragraph != "":
            content.append(template_p2.format(paragraph))
    content.append(template_p3)
    for paragraph in text_after_postanovlenie:
        if paragraph != "":
            content.append(template_p4.format(paragraph))
    today = date.today()
    content.append(template_p5.format(' '.join(author[0].split()), author[1], today.strftime('%Y'), today.strftime('%m'), os.path.basename(filepath)))
    print(' '.join(content))
    with open("../py/versions/output_file.txt", "w", encoding='utf-8') as file:
            # Write the content to the file
            file.write(' '.join(content))
        
root = Tk()     # создаем корневой объект - окно
root.title("Приложение на Tkinter")     # устанавливаем заголовок окна
root.geometry("300x250")    # устанавливаем размеры окна

lbl = ttk.Label(text="File Path:")
btn = ttk.Button(text="CLICK", command=open_file)


btn_create = ttk.Button(text="Создать", command=read_doc)
btn_create.pack(anchor=NW, padx=6, pady=6)
btn.pack()
lbl.pack()
root.mainloop()

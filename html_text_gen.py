import os
import sys
import re
from docx import Document
from datetime import date

template_p1 = '<p style="text-align: center;"><strong>Постановление </strong></p>\n<p style="text-align: right;"><strong>{}</strong></p>\n<p style="text-align: center;"><strong>{}</strong></p>\n'
template_p2 = '<p style="text-align: justify;">{}</p>\n'
template_p3 = '<p style="text-align: center;"><strong>П О С Т А Н О В Л Я Е Т:</strong></p>\n'
template_p4 = '<p style="text-align: justify;">{}</p>\n'
template_p5 = '<p style="text-align: center;">{}   <strong>{}</strong></p>\n<p style="text-align: right;"><strong><a href="http://pohr.ru/wps/wp-content/uploads/{}/{}/{}">Приложение</a></strong></p>\n'


def parse_document(filepath):
    num_lvl_counter = [0] * 9
    doc_date = ""
    html = []
    text_before_postanovlenie = []
    text_after_postanovlenie = []
    name = []
    if filepath != "":
        f = open(f'{filepath}', 'rb')
        document = Document(f)
        all_tables = document.tables
        all_paragraphs = document.paragraphs
        print('Всего таблиц в документе:', len(all_tables))
        data_tables = {i: None for i in range(len(all_tables))}
        # проходимся по таблицам
        for i, table in enumerate(all_tables):
            # print('\nДанные таблицы №', i)
            # создаем список строк для таблицы `i` (пока пустые)
            data_tables[i] = [[] for _ in range(len(table.rows))]
            # проходимся по строкам таблицы `i`
            for j, row in enumerate(table.rows):
                # проходимся по ячейкам таблицы `i` и строки `j`
                for cell in row.cells:
                    # добавляем значение ячейки в соответствующий
                    # список, созданного словаря под данные таблиц
                    data_tables[i][j].append(cell.text)
                    data_pattern = r'\d{2}\.\d{2}\.\d{4}\s*№\s*\d+'
                    match = re.search(data_pattern, cell.text)
                    if match:
                        print(match.group())
                        if doc_date == "":
                            doc_date = match.group()
                        break
        extracting = True
        lowest_num_lvl = 0
        for i, paragraph in enumerate(document.paragraphs):
            print(paragraph.text + "\n" + paragraph.style.name)
            style_name = paragraph.style.name if paragraph.style else "Normal"
            is_numbered = paragraph._element.pPr is not None and paragraph._element.pPr.numPr is not None
            num_pattern = r'\s*(\d+\.)+\s*'
            ilvl = -1
            match_in_ppr = False
            number_in_list = re.match(num_pattern, paragraph.text)
            if number_in_list:
                # print(f"re.findall = {re.findall(r'\d+\.', number_in_list.group())}")
                ilvl = len(re.findall(r'\d+\.', number_in_list.group())) - 1
                match_in_ppr = True
            if (is_numbered):
                ilvl = paragraph._element.pPr.numPr.ilvl.val if paragraph._element.pPr.numPr.ilvl is not None else 0
                # prefixes = ["1.", "a."]
                # prefix = prefixes[min(ilvl, len(prefixes)-1)]
            if ilvl != -1:
                num_lvl_counter[ilvl] += 1
                lowest_num_lvl = ilvl if ilvl > lowest_num_lvl else lowest_num_lvl
                if ilvl < lowest_num_lvl:
                    for i in range(ilvl + 1, len(num_lvl_counter)):
                        num_lvl_counter[i] = 0
                number = ".".join([str(i) for i in num_lvl_counter if i != 0])
                if match_in_ppr:
                    html.append(paragraph.text.strip())
                else:
                    html.append(number + ". " + paragraph.text.strip())
                print(f"Lvl: {ilvl}\nFinal number: {number}")
            else:
                html.append(paragraph.text.strip())
            # Вырезка сведений о постановлении
            if extracting == True:
                if paragraph.text != "":
                    text = paragraph.text
                    name.append(' '.join(text.split()))
                if name != [] and paragraph.text == "":
                    extracting = False
                else:
                    continue

        print(' '.join(name))
        print(name[len(name) - 1])
        start_before_ind = html.index(name[len(name) - 1]) if len(name) > 1 else html.index(name[len(name)])
        start_after_ind = 0
        name_pattern = r'([А-ЯЁ]\.\s*[А-ЯЁ]\.\s*[А-ЯЁ][а-яё]+)'
        for i in range(start_before_ind + 1, len(html)):
            if paragraph.text != "":
                continue
            else:
                if ''.join(html[i].lower().split()) != "ПОСТАНОВЛЯЕТ:".lower():
                    text_before_postanovlenie.append(' '.join(html[i].split()))
                else:
                    start_after_ind = i
                    break
        for i in range(start_after_ind + 1, len(html)):
            if paragraph.text != "":
                continue
            else:
                if html[i - 1] == "" and re.search(name_pattern, html[i]) and len(
                        html[i]) < 150:  # Если находит строку с фамилией автора в коротком абзаце
                    global author
                    author = re.split(name_pattern, html[i])
                    print(author)
                    break
                else:
                    text_after_postanovlenie.append(' '.join(html[i].split()))
        content = []
        content.append(template_p1.format(doc_date, ' '.join(name)))
        for paragraph in text_before_postanovlenie:
            if paragraph != "":
                content.append(template_p2.format(paragraph))
        content.append(template_p3)
        for paragraph in text_after_postanovlenie:
            if paragraph != "":
                content.append(template_p4.format(paragraph))
        today = date.today()
        content.append(
            template_p5.format(' '.join(author[0].split()), author[1], today.strftime('%Y'), today.strftime('%m'),
                               os.path.basename(filepath)))
        print(' '.join(content))
        return ' '.join(content)



"""
if len(sys.argv) != 2:
    print("Использование: python script.py <input_file> <output_file>")
    sys.exit(1)
else:
    input_file = sys.argv[1]
    parse_document(input_file)"""
